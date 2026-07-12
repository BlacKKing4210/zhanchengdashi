#!/usr/bin/env python3
"""Build the editable Word review copy for the game audio design."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "GAME_AUDIO_DESIGN.md"
OUTPUT = ROOT / "docs" / "GAME_AUDIO_DESIGN.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "17365D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "667085"
INK = "202124"


def set_run_font(run, size: float = 11, bold: bool = False, color: str = INK, italic: bool = False) -> None:
	run.font.name = "Calibri"
	run.font.size = Pt(size)
	run.font.bold = bold
	run.font.italic = italic
	run.font.color.rgb = RGBColor.from_string(color)
	run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
	run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
	run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_repeat_table_header(row) -> None:
	tr_pr = row._tr.get_or_add_trPr()
	tbl_header = OxmlElement("w:tblHeader")
	tbl_header.set(qn("w:val"), "true")
	tr_pr.append(tbl_header)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
	tc = cell._tc
	tc_pr = tc.get_or_add_tcPr()
	tc_mar = tc_pr.first_child_found_in("w:tcMar")
	if tc_mar is None:
		tc_mar = OxmlElement("w:tcMar")
		tc_pr.append(tc_mar)
	for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
		node = tc_mar.find(qn(f"w:{tag}"))
		if node is None:
			node = OxmlElement(f"w:{tag}")
			tc_mar.append(node)
		node.set(qn("w:w"), str(value))
		node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
	tc_pr = cell._tc.get_or_add_tcPr()
	shd = tc_pr.find(qn("w:shd"))
	if shd is None:
		shd = OxmlElement("w:shd")
		tc_pr.append(shd)
	shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int]) -> None:
	table.alignment = WD_TABLE_ALIGNMENT.LEFT
	table.autofit = False
	tbl_pr = table._tbl.tblPr
	tbl_w = tbl_pr.find(qn("w:tblW"))
	if tbl_w is None:
		tbl_w = OxmlElement("w:tblW")
		tbl_pr.append(tbl_w)
	tbl_w.set(qn("w:w"), str(sum(widths)))
	tbl_w.set(qn("w:type"), "dxa")
	tbl_ind = tbl_pr.find(qn("w:tblInd"))
	if tbl_ind is None:
		tbl_ind = OxmlElement("w:tblInd")
		tbl_pr.append(tbl_ind)
	tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
	tbl_ind.set(qn("w:type"), "dxa")
	tbl_layout = tbl_pr.find(qn("w:tblLayout"))
	if tbl_layout is None:
		tbl_layout = OxmlElement("w:tblLayout")
		tbl_pr.append(tbl_layout)
	tbl_layout.set(qn("w:type"), "fixed")
	grid = table._tbl.tblGrid
	for child in list(grid):
		grid.remove(child)
	for width in widths:
		grid_col = OxmlElement("w:gridCol")
		grid_col.set(qn("w:w"), str(width))
		grid.append(grid_col)
	for row in table.rows:
		for index, cell in enumerate(row.cells):
			width = widths[index]
			tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
			if tc_w is None:
				tc_w = OxmlElement("w:tcW")
				cell._tc.get_or_add_tcPr().append(tc_w)
			tc_w.set(qn("w:w"), str(width))
			tc_w.set(qn("w:type"), "dxa")
			set_cell_margins(cell)
			cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_inline_runs(paragraph, text: str, size: float = 11, color: str = INK) -> None:
	parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
	for part in parts:
		if not part:
			continue
		if part.startswith("`") and part.endswith("`"):
			run = paragraph.add_run(part[1:-1])
			set_run_font(run, size=size, bold=True, color=DARK_BLUE)
		elif part.startswith("**") and part.endswith("**"):
			run = paragraph.add_run(part[2:-2])
			set_run_font(run, size=size, bold=True, color=color)
		else:
			run = paragraph.add_run(part)
			set_run_font(run, size=size, color=color)


def configure_styles(doc: Document) -> None:
	styles = doc.styles
	normal = styles["Normal"]
	normal.font.name = "Calibri"
	normal.font.size = Pt(11)
	normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
	normal.paragraph_format.space_before = Pt(0)
	normal.paragraph_format.space_after = Pt(6)
	normal.paragraph_format.line_spacing = 1.25
	for name, size, before, after, color in (
		("Heading 1", 16, 18, 10, BLUE),
		("Heading 2", 13, 14, 7, BLUE),
		("Heading 3", 12, 10, 5, DARK_BLUE),
	):
		style = styles[name]
		style.font.name = "Calibri"
		style.font.size = Pt(size)
		style.font.bold = True
		style.font.color.rgb = RGBColor.from_string(color)
		style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
		style.paragraph_format.space_before = Pt(before)
		style.paragraph_format.space_after = Pt(after)
		style.paragraph_format.keep_with_next = True
	for name in ("List Bullet", "List Number"):
		style = styles[name]
		style.font.name = "Calibri"
		style.font.size = Pt(11)
		style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
		style.paragraph_format.left_indent = Inches(0.375)
		style.paragraph_format.first_line_indent = Inches(-0.188)
		style.paragraph_format.space_after = Pt(4)
		style.paragraph_format.line_spacing = 1.25


def add_page_number(paragraph) -> None:
	paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	run = paragraph.add_run("第 ")
	set_run_font(run, size=9, color=MUTED)
	fld_begin = OxmlElement("w:fldChar")
	fld_begin.set(qn("w:fldCharType"), "begin")
	instr = OxmlElement("w:instrText")
	instr.set(qn("xml:space"), "preserve")
	instr.text = " PAGE "
	fld_sep = OxmlElement("w:fldChar")
	fld_sep.set(qn("w:fldCharType"), "separate")
	fld_text = OxmlElement("w:t")
	fld_text.text = "1"
	fld_end = OxmlElement("w:fldChar")
	fld_end.set(qn("w:fldCharType"), "end")
	for node in (fld_begin, instr, fld_sep, fld_text, fld_end):
		run._r.append(node)
	run2 = paragraph.add_run(" 页")
	set_run_font(run2, size=9, color=MUTED)


def add_masthead(doc: Document) -> None:
	section = doc.sections[0]
	header = section.header.paragraphs[0]
	header.alignment = WD_ALIGN_PARAGRAPH.LEFT
	run = header.add_run("丛林法则  |  游戏音频系统规格")
	set_run_font(run, size=9, bold=True, color=MUTED)
	add_page_number(section.footer.paragraphs[0])
	spacer = doc.add_paragraph()
	spacer.paragraph_format.space_after = Pt(12)
	title = doc.add_paragraph()
	title.paragraph_format.space_after = Pt(4)
	run = title.add_run("全游戏音乐与音效设计")
	set_run_font(run, size=24, bold=True, color=DARK_BLUE)
	subtitle = doc.add_paragraph()
	subtitle.paragraph_format.space_after = Pt(12)
	add_inline_runs(subtitle, "全局音频方向、资源清单、Godot 运行时契约与 QA 标准", size=12, color=MUTED)
	meta = doc.add_paragraph()
	meta.paragraph_format.space_after = Pt(12)
	add_inline_runs(meta, "版本 v1.0  |  状态：已确认实装  |  2026-07-12", size=10, color=MUTED)
	rule = doc.add_paragraph()
	rule.paragraph_format.space_after = Pt(8)
	p_pr = rule._p.get_or_add_pPr()
	p_bdr = OxmlElement("w:pBdr")
	bottom = OxmlElement("w:bottom")
	bottom.set(qn("w:val"), "single")
	bottom.set(qn("w:sz"), "10")
	bottom.set(qn("w:space"), "1")
	bottom.set(qn("w:color"), BLUE)
	p_bdr.append(bottom)
	p_pr.append(p_bdr)


def table_widths(column_count: int) -> list[int]:
	if column_count == 4:
		return [1300, 1800, 3160, 3100]
	if column_count == 3:
		return [1700, 2300, 5360]
	if column_count == 2:
		return [2500, 6860]
	base = CONTENT_WIDTH_DXA // column_count
	widths = [base] * column_count
	widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
	return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
	column_count = len(rows[0])
	table = doc.add_table(rows=len(rows), cols=column_count)
	table.style = "Table Grid"
	set_table_geometry(table, table_widths(column_count))
	set_repeat_table_header(table.rows[0])
	for row_index, values in enumerate(rows):
		for column_index, value in enumerate(values):
			cell = table.cell(row_index, column_index)
			cell.text = ""
			paragraph = cell.paragraphs[0]
			paragraph.paragraph_format.space_before = Pt(0)
			paragraph.paragraph_format.space_after = Pt(0)
			paragraph.paragraph_format.line_spacing = 1.15
			add_inline_runs(paragraph, value, size=9.5 if column_count >= 3 else 10, color=INK)
			if row_index == 0:
				set_cell_shading(cell, LIGHT_BLUE)
				for run in paragraph.runs:
					run.bold = True
			elif row_index % 2 == 0:
				set_cell_shading(cell, LIGHT_GRAY)
	spacer = doc.add_paragraph()
	spacer.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
	rows: list[list[str]] = []
	index = start
	while index < len(lines) and lines[index].strip().startswith("|"):
		values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
		if not all(re.fullmatch(r":?-{3,}:?", value.replace(" ", "")) for value in values):
			rows.append(values)
		index += 1
	return rows, index


def build_document() -> None:
	lines = SOURCE.read_text(encoding="utf-8").splitlines()
	doc = Document()
	section = doc.sections[0]
	section.start_type = WD_SECTION.NEW_PAGE
	section.top_margin = Inches(1)
	section.bottom_margin = Inches(1)
	section.left_margin = Inches(1)
	section.right_margin = Inches(1)
	section.header_distance = Inches(0.492)
	section.footer_distance = Inches(0.492)
	configure_styles(doc)
	add_masthead(doc)

	index = 0
	while index < len(lines):
		line = lines[index].strip()
		if not line or line.startswith("# "):
			index += 1
			continue
		if line.startswith("> "):
			paragraph = doc.add_paragraph()
			paragraph.paragraph_format.left_indent = Inches(0.16)
			paragraph.paragraph_format.right_indent = Inches(0.1)
			paragraph.paragraph_format.space_before = Pt(4)
			paragraph.paragraph_format.space_after = Pt(8)
			set_cell_shading_like_paragraph(paragraph, "EDF4FB")
			add_inline_runs(paragraph, line[2:], size=10.5, color=DARK_BLUE)
			index += 1
			continue
		if line.startswith("### "):
			paragraph = doc.add_paragraph(style="Heading 2")
			add_inline_runs(paragraph, line[4:], size=13, color=BLUE)
			index += 1
			continue
		if line.startswith("## "):
			paragraph = doc.add_paragraph(style="Heading 1")
			add_inline_runs(paragraph, line[3:], size=16, color=BLUE)
			index += 1
			continue
		if line.startswith("|"):
			rows, index = parse_table(lines, index)
			if rows:
				add_markdown_table(doc, rows)
			continue
		if line.startswith("- "):
			paragraph = doc.add_paragraph(style="List Bullet")
			add_inline_runs(paragraph, line[2:])
			index += 1
			continue
		if re.match(r"^\d+\.\s", line):
			paragraph = doc.add_paragraph(style="List Number")
			add_inline_runs(paragraph, re.sub(r"^\d+\.\s+", "", line))
			index += 1
			continue
		paragraph = doc.add_paragraph()
		add_inline_runs(paragraph, line)
		index += 1

	properties = doc.core_properties
	properties.title = "全游戏音乐与音效设计"
	properties.subject = "丛林法则全局音频系统"
	properties.author = "Codex Game Studio"
	properties.keywords = "Godot, 游戏音频, BGM, SFX"
	doc.save(OUTPUT)
	print(OUTPUT)


def set_cell_shading_like_paragraph(paragraph, fill: str) -> None:
	p_pr = paragraph._p.get_or_add_pPr()
	shd = OxmlElement("w:shd")
	shd.set(qn("w:fill"), fill)
	p_pr.append(shd)


if __name__ == "__main__":
	build_document()
