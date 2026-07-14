from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PLAYER_ACCOUNT_AND_SERVER_PROFILE_DESIGN.docx"


def set_font(run, size=11, bold=False, color="20252B"):
    run.font.name = "Microsoft YaHei"
    fonts = run._element.get_or_add_rPr().rFonts
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for name, size, before, after in [("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6)]:
    style = doc.styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(46, 116, 181)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

list_style = doc.styles["List Bullet"]
list_style.font.name = "Microsoft YaHei"
list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
list_style.paragraph_format.left_indent = Inches(0.5)
list_style.paragraph_format.first_line_indent = Inches(-0.25)
list_style.paragraph_format.space_after = Pt(8)
list_style.paragraph_format.line_spacing = 1.167

header = section.header.paragraphs[0]
header.text = "战城大师 · 功能规格"
set_font(header.runs[0], 9, color="6B7280")
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.text = "v1.1 · 2026-07-14"
set_font(footer.runs[0], 9, color="6B7280")

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
set_font(title.add_run("玩家账号与服务器资料设计"), 23, True, "111827")
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
set_font(subtitle.add_run("安装 ID 自动登录、服务端权威进度与基地账号中心"), 13, color="4B5563")

sections = [
    ("1. 目标", [
        "每名玩家拥有服务器生成且不可修改的 UserID；卡牌数量、卡牌星级、出战编组、抽卡券、段位星数与 ELO 均以服务器资料为权威。",
        "首次联网自动创建账号并默认登录，后续启动继续使用同一账号，不依赖 MAC、硬盘序列号或其他物理地址。",
    ]),
    ("2. 默认设备账号流程", [
        "客户端首次启动生成 32 字节随机安装 ID，并保存到 user://client/device_account.json。",
        "首次连接时服务器创建游客账号和 UserID，仅在首次响应中签发 32 字节随机长期登录令牌。",
        "后续连接使用安装 ID 与长期令牌自动登录；服务器再签发仅用于当前连接的短期会话令牌。",
        "注销只结束当前会话，不删除本机长期令牌；下次启动或重新连接仍默认进入原账号。",
    ]),
    ("3. 安全边界", [
        "服务器只保存安装 ID 的 SHA-256 摘要、随机盐和长期令牌哈希，不保存安装 ID 或令牌明文。",
        "客户端只在 user:// 保存本机凭据，不把安装 ID 当作密码；长期令牌错误时服务器统一拒绝认证。",
        "网络 peer 只能读写其当前会话对应的玩家资料；断线与注销都会解除 peer 到用户的映射。",
        "删除应用数据或丢失本机凭据会失去自动登录能力；正式发行前需补充账号绑定、令牌轮换和跨设备找回。",
    ]),
    ("4. 兼容功能", [
        "原账号与密码注册/登录入口继续保留，作为兼容入口和未来账号绑定基础；密码只保存加盐派生哈希。",
        "主页面点击基地打开账号中心，可查看 UserID、注销当前会话、阅读玩家协议并切换音乐与音效。",
    ]),
    ("5. 验收", [
        "全新安装首次连接后自动得到 UserID、长期令牌和默认服务器资料。",
        "服务器重启后，同一安装 ID 与长期令牌登录到相同 UserID，且长期令牌不会再次返回。",
        "不同安装获得不同 UserID；伪造安装 ID、错误令牌及未认证资料写入均被拒绝。",
        "账号存储测试、ENet 双客户端自动登录回归和 Godot 全项目解析通过。",
    ]),
]

for heading, items in sections:
    doc.add_heading(heading, level=1)
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        set_font(paragraph.add_run(item))

doc.core_properties.title = "玩家账号与服务器资料设计"
doc.core_properties.subject = "安装 ID 自动登录与服务端权威资料规格"
doc.core_properties.author = "Codex Game Studio"
doc.save(OUT)
print(OUT)
