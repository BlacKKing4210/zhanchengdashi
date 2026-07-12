#!/usr/bin/env python3
"""Build the user-facing Word review document for universal animal motion FX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "ANIMAL_UNIVERSAL_MOTION_FEEDBACK_REVIEW.docx"
VISUALS = ROOT / "output" / "visual_concepts"

INK = "24211D"
PAPER = "F7E6B5"
HIT = "D9543D"
POWER = "F2C14E"
GAIN = "2D8C7A"
MOTION = "4E9BC4"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
MUTED = "666666"
TABLE_HEADER = "E8EEF5"
WHITE = "FFFFFF"

# Preset: compact_reference_guide.
# Named override: all East Asian glyphs use Microsoft YaHei for CJK coverage;
# Latin glyphs retain the preset's Calibri family.


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, **kwargs) -> None:
    for run in paragraph.runs:
        set_run_font(run, **kwargs)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_bottom_border(paragraph, color: str = INK, size: int = 10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Twips(width_dxa)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths_dxa)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def format_table(table, header: bool = True) -> None:
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=9.3, color=INK, bold=(header and row_index == 0))
            if header and row_index == 0:
                set_cell_shading(cell, TABLE_HEADER)
    if header:
        mark_header_row(table.rows[0])


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED)


def add_figure(doc: Document, path: Path, width_inches: float, alt_text: str, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=Inches(width_inches))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", caption)
    add_caption(doc, caption)


def add_callout(doc: Document, label: str, text: str, fill: str = "FFF4D4") -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.10)
    paragraph.paragraph_format.right_indent = Inches(0.06)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), POWER)
    borders.append(left)
    p_pr.append(borders)
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, size=11, bold=True, color=INK)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=11, color=INK)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, bold=True, color=BLUE_DARK)
    value_run = paragraph.add_run(text)
    set_run_font(value_run, color=INK)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)
    set_run_font(run, size=9, color=MUTED)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = paragraph.add_run("动物通用动作反馈")
    set_run_font(left, size=9, bold=True, color=MUTED)
    right = paragraph.add_run("\t效果图评审稿")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.clear()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer_paragraph.add_run("占城大师  |  ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(footer_paragraph)


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(4)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("视觉表现设计 / 评审稿")
    set_run_font(run, size=10, bold=True, color=MOTION)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("动物通用动作反馈")
    set_run_font(title_run, size=26, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run("以程序化姿态与通用图片 FX 替代逐动物序列帧")
    set_run_font(subtitle_run, size=13, color=BLUE_DARK)

    metadata = [
        ("状态", "纯效果图评审；尚未进入 Godot 实装"),
        ("推荐", "A 弹性墨线动作"),
        ("基线", "当前动物 PNG 不变；720 x 1280；7 x 13；战斗动物 44 px"),
        ("日期", "2026-07-12"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, size=10.5, bold=True, color=INK)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(10)
    add_bottom_border(rule, color=INK, size=10)


def add_motion_table(doc: Document) -> None:
    rows = [
        ("移动", "上下弹跳、左右倾斜、脚点压缩；方向来自实际位移", "0.32s", "速度线 + 脚边尘点"),
        ("攻击", "反向蓄力 -> 朝目标前冲 8-10 px -> 回弹", "0.24s", "方向攻击弧 + 接触火花"),
        ("受击", "当帧闪白、横向挤压、沿来源反向视觉后退 4-6 px", "0.16s", "珊瑚红冲击缺口 + 奶白碎片"),
        ("获得属性", "图片轻弹；同类事件短时间合并，持续光环不重复触发", "0.55s", "上升环 + 属性小图形"),
        ("获得提升", "1.00 -> 1.16 -> 1.00；只用于重要提升确认", "0.68s", "金色星环 + 上升刻度"),
        ("死亡", "保存图片快照后倾倒、缩小、淡出；清除附着状态", "0.42s", "断裂墨圈 + 尘烟 + 消散碎点"),
    ]
    table = doc.add_table(rows=1, cols=4)
    headers = ["状态", "动作拟态", "时长", "图片反馈"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    set_table_geometry(table, [1150, 4300, 950, 2960])
    format_table(table)


def add_option_page(doc: Document, option: str, title: str, summary: str, risk: str, verdict: str) -> None:
    doc.add_page_break()
    doc.add_heading(f"方案 {option.upper()} - {title}", level=1)
    add_labeled_paragraph(doc, "定位", summary)
    add_labeled_paragraph(doc, "风险", risk)
    add_labeled_paragraph(doc, "结论", verdict)
    add_figure(
        doc,
        VISUALS / f"animal_universal_motion_option_{option}_review_board.png",
        6.35,
        f"方案 {option.upper()} 的攻击、受击、移动、属性、提升和死亡三阶段动作评审板",
        f"图 {option.upper()}：{title}，六类动作三阶段与实际 44 px 预览",
    )


def build_docx() -> None:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "动物通用动作反馈设计"
    doc.core_properties.subject = "通用动物 UI 动效效果图评审与 Godot 落地规格"
    doc.core_properties.author = "占城大师项目组"

    add_title_block(doc)
    add_callout(
        doc,
        "推荐决策",
        "选择 A 弹性墨线动作。它在 44 px 动物尺寸下仍能区分六类状态，且比 B 更少遮挡、比 C 更强可读。用户明确确认并说“实装”前，不修改运行时。",
    )
    add_figure(
        doc,
        VISUALS / "animal_universal_motion_options_overview.png",
        6.35,
        "A、B、C 三套动物通用动作反馈在当前 720 x 1280 与 7 x 13 战斗几何中的并排对比",
        "图 1：三版实际战场尺寸对比；动物原图、格位与血条层保持不变",
    )

    doc.add_page_break()
    doc.add_heading("1. 目标与工程基线", level=1)
    add_labeled_paragraph(doc, "目标", "用一套通用动作系统弥补动物只有单张 PNG 的缺陷，不为 60 个动物分别制作序列帧。")
    add_labeled_paragraph(doc, "不变项", "动物图片、白边、卡牌绑定、逻辑坐标、碰撞、攻击范围、棋盘、点击区和战斗数值。")
    add_labeled_paragraph(doc, "工程基线", "720 x 1280 设计画布；592 x 982 经典战斗内场；7 x 13 共 91 格；HEX_SIZE 43；战斗动物 44 x 44。")
    add_labeled_paragraph(doc, "实现原则", "脚点程序化缩放/旋转/绘制位移 + 通用图片 FX；图片只表达动作，不承担逻辑。")
    doc.add_heading("2. 六类动作规格", level=1)
    add_motion_table(doc)
    add_callout(doc, "建议补充", "出生反馈使用 0.25s 脚点弹出和短尘圈，形成完整生命周期，并与金色“获得提升”语义分离。", fill="EAF3F1")

    add_option_page(
        doc,
        "a",
        "弹性墨线动作",
        "中等粗细墨线、清楚接触帧、形状数量适中，动作和遮挡最平衡。",
        "强度控制不当时攻击弧仍可能侵入相邻格，因此运行时需限定最大半径。",
        "当前推荐默认套件；Boss 或关键技能只增加一次强度，不另建语言。",
    )
    add_option_page(
        doc,
        "b",
        "粗笔战斗印章",
        "更粗轮廓和更大的环形缺口，在小屏与复杂背景上识别最快。",
        "多人同屏时最容易遮挡相邻动物、伤害数字和血条。",
        "保留为“强反馈”档，不建议所有普通事件默认使用。",
    )
    add_option_page(
        doc,
        "c",
        "克制舞台轨迹",
        "更少形状、更低覆盖面积，长期观看最安静。",
        "浅色地块与大量单位同时移动时，攻击/移动识别略弱。",
        "适合作为减弱动效模式的视觉基线。",
    )

    doc.add_page_break()
    doc.add_heading("3. 语义色、图层与锚点", level=1)
    token_rows = [
        ("fx_ink", "#24211D", "轮廓、死亡断环、速度线", "断裂/方向"),
        ("fx_paper", "#F7E6B5", "接触闪光、碎片、星点", "四角星/碎片"),
        ("fx_hit", "#D9543D", "受击、危险接触", "锯齿缺口"),
        ("fx_power", "#F2C14E", "攻击接触、重要提升", "星环/上升刻度"),
        ("fx_gain", "#2D8C7A", "属性获得、恢复、护盾", "上升箭头/环"),
        ("fx_motion", "#4E9BC4", "移动、速度属性", "平行轨迹"),
    ]
    table = doc.add_table(rows=1, cols=4)
    for index, value in enumerate(("Token", "初始色", "用途", "非颜色线索")):
        table.rows[0].cells[index].text = value
    for row in token_rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    set_table_geometry(table, [1800, 1500, 3400, 2660])
    format_table(table)
    add_labeled_paragraph(doc, "颜色规则", "禁止渐变、外发光与多层阴影；颜色不能单独承担语义。")
    add_labeled_paragraph(doc, "统一锚点", "动物图片底部中心，约等于当前绘制位置 pos + Vector2(0, 14)。")
    add_labeled_paragraph(doc, "绘制顺序", "地块/建筑 -> 背后轨迹与成长环 -> 动物 PNG -> 接触前景 -> 血条/队伍标记/伤害数字 -> 全局 UI。")

    doc.add_heading("4. Godot 落地计划（获批后）", level=1)
    add_labeled_paragraph(doc, "状态索引", "使用稳定 unit[\"id\"] 管理 unit_motion_states，不能使用会随死亡过滤变化的数组下标。")
    add_labeled_paragraph(doc, "附着状态", "moving、move_phase、facing、attack_time/direction、hit_time/direction。")
    add_labeled_paragraph(doc, "脱离效果", "impact、stat_gain、power_up、death_snapshot、dust 继续进入 effects。")
    add_labeled_paragraph(doc, "事件挂点", "_update_units()、_unit_attack_target()、_damage_unit()、属性增加函数、_try_upgrade_selected_card()、_handle_unit_death()、_spawn_unit()。")
    add_labeled_paragraph(doc, "优先级", "death_snapshot > hit > attack > move > idle。")
    add_labeled_paragraph(doc, "架构取舍", "维持当前 Node2D 即时绘制，不为本功能重构全部动物为 Sprite2D。")

    doc.add_page_break()
    doc.add_heading("5. 性能、减弱动效与 QA Gate", level=1)
    add_labeled_paragraph(doc, "性能", "同屏短效建议上限 32，对象池容量 64；超限先丢移动尘点，再合并同类属性反馈；离屏 FX 不绘制。")
    add_labeled_paragraph(doc, "减弱动效", "关闭旋转、前冲、后退和循环弹跳，仅保留 0.12-0.25s 静态符号、一次轻微缩放与淡出。")
    add_labeled_paragraph(doc, "中断", "动画必须可中断、可复位，不锁输入；死亡清除同单位所有附着状态。")

    qa = doc.add_table(rows=1, cols=3)
    for index, value in enumerate(("级别", "必须满足", "当前提案结果")):
        qa.rows[0].cells[index].text = value
    qa_rows = [
        ("P0", "不改动物/几何/逻辑；六类语义独立；死亡快照可见；结束后复位", "静态效果图通过"),
        ("P1", "44 px 可读；不遮血条与伤害数字；多体型/并发/连续受击通过", "A 通过；B 遮挡风险最高；C 强度偏低"),
        ("P2", "三版强度差异清楚；减弱模式仍有反馈；边界分辨率截图对照", "待获批实装后完成动态验证"),
    ]
    for row in qa_rows:
        cells = qa.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    set_table_geometry(qa, [1000, 5200, 3160])
    format_table(qa)
    add_callout(
        doc,
        "下一步",
        "请确认 A、B 或 C。若选择 A 并明确说“实装”，再进入 atlas 导入、运行时状态、事件触发、减弱动效、目标/边界分辨率截图和 Godot 回归测试。",
    )

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
