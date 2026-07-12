#!/usr/bin/env python3
"""Build the editable Word specification for multiplayer elimination transfer rules."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import build_game_audio_design_docx as base


ROOT = Path(__file__).resolve().parents[1]


def add_masthead(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("丛林法则  |  多人战斗系统规格")
    base.set_run_font(run, size=9, bold=True, color=base.MUTED)
    base.add_page_number(section.footer.paragraphs[0])

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("多人淘汰后的领地与大本营转移规则")
    base.set_run_font(run, size=23, bold=True, color=base.DARK_BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    base.add_inline_runs(subtitle, "领地锁定重置、俘获基地功能、重复易主与多人结算边界", size=12, color=base.MUTED)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    base.add_inline_runs(meta, "版本 v1.0  |  状态：已确认实装  |  2026-07-12", size=10, color=base.MUTED)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), base.BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def main() -> None:
    base.SOURCE = ROOT / "docs" / "MULTIPLAYER_ELIMINATION_TRANSFER_RULE.md"
    base.OUTPUT = ROOT / "docs" / "MULTIPLAYER_ELIMINATION_TRANSFER_RULE.docx"
    base.add_masthead = add_masthead
    base.build_document()
    print(base.OUTPUT)


if __name__ == "__main__":
    main()
