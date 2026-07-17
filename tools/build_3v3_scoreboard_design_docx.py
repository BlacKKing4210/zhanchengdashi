from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "THREE_VS_THREE_SCOREBOARD_DESIGN.docx"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 100, 112)


def set_run_font(run, size, color=INK, bold=False):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def set_paragraph_format(paragraph, before=0, after=6, line=1.1):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bottom_border(paragraph, color="2E74B5"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "10")
    border.set(qn("w:space"), "6")
    border.set(qn("w:color"), color)
    borders.append(border)
    p_pr.append(borders)


def configure_style(style, font_name, size, color, before, after, line, bold=False):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def add_section(doc, title, body):
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run(title)
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, after=7, line=1.1)
    run = paragraph.add_run(body)
    set_run_font(run, 11, RGBColor(30, 30, 30))


def add_definition(doc, label, body):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, after=5, line=1.1)
    label_run = paragraph.add_run(label + "：")
    set_run_font(label_run, 11, INK, True)
    body_run = paragraph.add_run(body)
    set_run_font(body_run, 11, RGBColor(30, 30, 30))


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(doc.styles["Normal"], "Microsoft YaHei", 11, RGBColor(30, 30, 30), 0, 6, 1.1)
    configure_style(doc.styles["Heading 1"], "Microsoft YaHei", 16, BLUE, 16, 8, 1.1, True)
    configure_style(doc.styles["Heading 2"], "Microsoft YaHei", 13, BLUE, 12, 6, 1.1, True)
    configure_style(doc.styles["Heading 3"], "Microsoft YaHei", 12, INK, 8, 4, 1.1, True)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(header, after=0)
    set_run_font(header.add_run("丛林法则 | 战斗 HUD 功能规格"), 9, MUTED, True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(footer, before=0, after=0)
    set_run_font(footer.add_run("内部规格 · 2026-07-18"), 9, MUTED)

    title = doc.add_paragraph()
    set_paragraph_format(title, after=4)
    set_run_font(title.add_run("3V3 阵营地块记分牌"), 23, INK, True)
    subtitle = doc.add_paragraph()
    set_paragraph_format(subtitle, after=14)
    set_run_font(subtitle.add_run("双边房间战斗 HUD · 显示规则、数据口径与验收标准"), 12, MUTED)
    rule = doc.add_paragraph()
    set_paragraph_format(rule, after=8)
    add_bottom_border(rule)

    add_section(
        doc,
        "目标",
        "在双边 3V3 房间战斗中持续显示双方当前地块控制数，让玩家不离开棋盘即可比较阵营态势与实际颜色归属。",
    )
    add_section(
        doc,
        "HUD 规则",
        "仅在多人、非六人自由混战、且每边三名玩家时显示。顶部记分牌固定为 A 方（槽位 1-3）与 B 方（槽位 4-6）两栏；每栏包含阵营名、三个槽位色点和“地块数 + 格”。本地所在阵营额外使用金色描边和“我方”标签，对方显示“敌方”。",
    )
    add_section(
        doc,
        "数据口径",
        "地块数按 BoardRules.visual_owner(tile) 聚合，不只统计已真实解锁的 tile.team。因此预归属、软占领、真实解锁和基地易主后的可见阵营色与记分牌数保持一致。现有联网快照已同步完整 tiles，功能不新增服务端字段或网络消息。",
    )
    add_section(
        doc,
        "实现边界",
        "main.gd 只负责显示条件、阵营聚合与绘制；测试覆盖初始地块数、软占领后的实时变化、颜色槽位与模式门槛。不修改结算、地块归属、盟友关系、自由混战排行榜或服务端快照协议。",
    )
    add_section(doc, "验收标准", "以下条件全部成立后视为完成。")
    add_definition(doc, "可见性", "任意 3V3 地图顶部同时显示 A、B 两栏；1V1、2V2 和自由混战不显示。")
    add_definition(doc, "准确性", "每栏数值等于该阵营所有槽位的可见归属格总数，并在占领、转移或基地易主后的下一帧更新。")
    add_definition(doc, "颜色", "每栏三个色点与本局棋盘中对应槽位的阵营色一致；我方栏始终有金色强调。")
    add_definition(doc, "回归", "多人匹配规则、自由混战排行榜和经典战斗回归测试保持通过。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
