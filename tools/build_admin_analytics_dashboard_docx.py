from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "ADMIN_ANALYTICS_DASHBOARD_DESIGN.docx"

BLUE = RGBColor(46, 116, 181)
INK = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 102, 120)
LIGHT = "E8EEF5"


def _set_font(run, size: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _set_cell_fill(cell, value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), value)
    tc_pr.append(shade)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _fixed_table(table, widths_inches: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[index])
            _set_cell_margins(cell)


def _paragraph(doc, text: str = "", style: str | None = None, bold=False, color=INK, size=11):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    _set_font(run, size, color, bold)
    return paragraph


def _bullet(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    _set_font(paragraph.add_run(text), 11, INK)


def _heading(doc, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10}[level])
    paragraph.paragraph_format.space_after = Pt({1: 10, 2: 7, 3: 5}[level])
    run = paragraph.add_run(text)
    _set_font(run, {1: 16, 2: 13, 3: 12}[level], BLUE if level < 3 else INK, True)


def _table(doc, rows: list[tuple[str, str]], widths=(1.875, 4.625)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _fixed_table(table, list(widths))
    headers = ("项目", "规则")
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_fill(cell, LIGHT)
        run = cell.paragraphs[0].add_run(text)
        _set_font(run, 10.5, INK, True)
    for left, right in rows:
        cells = table.add_row().cells
        for index, text in enumerate((left, right)):
            run = cells[index].paragraphs[0].add_run(text)
            _set_font(run, 10.5, INK, index == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(header.add_run("Jungle Law | 运营数据后台"), 9, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(footer.add_run("内部使用 | 2026-07-18"), 9, MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _set_font(title.add_run("竞技数据统计后台"), 23, INK, True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_font(subtitle.add_run("设计说明与安全实施契约"), 13, MUTED)
    _table(doc, [
        ("任务", "OPS-ANALYTICS-001"),
        ("状态", "实施中"),
        ("目标", "授权运营人员查看排行榜、头部卡组与动物胜率"),
        ("数据来源", "服务器持久化的脱敏统计快照"),
    ])

    _heading(doc, "目标与范围", 1)
    _paragraph(doc, "建立一个仅限所有者及被授权运营人员访问的网页后台，以原创赛事积分榜的信息层级展示排行榜、卡组、动物胜率和最近对局。界面借鉴国际足球赛事的数据组织方式，但不使用第三方赛事名称、标志或资产。")
    _bullet(doc, "Dedicated server 独占统计投影和持久化。")
    _bullet(doc, "独立 Node.js 网站仅读取脱敏快照，绝不读取原始玩家账号库。")
    _bullet(doc, "Owner 可授权与撤销 Analyst；Analyst 仅可查看运营数据。")
    _bullet(doc, "默认绑定 127.0.0.1:24568；远程访问必须启用 TLS。")

    _heading(doc, "页面与交互", 1)
    _table(doc, [
        ("赛事总览", "KPI、积分榜式排行榜、近期对局、数据来源与更新时间。"),
        ("玩家榜与卡组", "按段位、星级、ELO、胜率、场次排序；展开查看已冻结卡组和等级。"),
        ("动物胜率", "展示携带场次、胜负、胜率和样本量；样本为 0 时显示暂无数据。"),
        ("授权与审计", "Owner 管理人员并审查登录、失败登录、授权和禁用记录。"),
    ])

    _heading(doc, "数据契约", 1)
    _paragraph(doc, "服务器在对局开始时冻结参赛者、队伍、卡组和卡牌等级。在首次收到同一 match_id 的终局快照时只结算一次。网页仅消费 dashboard_snapshot.json，其中包含 overview、leaderboard、top_decks、animals 和 recent_matches。")
    _paragraph(doc, "动物胜率的定义为“携带该动物的已完成对局胜率”，不是伤害占比或最后一击概率。每个数值必须同时显示样本量。", bold=True)

    _heading(doc, "安全与权限", 1)
    _table(doc, [
        ("初始 Owner", "只能在服务器本机 CLI 初始化；不存在默认账号或密码。"),
        ("密码", "Node crypto.scrypt、独立随机盐、常数时间比较；不复用玩家密码。"),
        ("会话", "HttpOnly、SameSite=Strict；远程 HTTPS 模式强制 Secure。"),
        ("远程访问", "非回环绑定只有同时提供 TLS 证书和私钥时才允许。"),
        ("写操作", "CSRF 令牌、同源校验、审计日志与二次确认。"),
    ])

    _heading(doc, "已知限制", 1)
    _paragraph(doc, "当前多人战斗由房主客户端模拟，服务器持久化的是已验证 match_id、冻结参赛名单及房主权威终局快照。因此本版本标记为服务器记录的运营统计；若需要反作弊官方战绩，下一阶段需要将战斗模拟和结算迁入 dedicated server。")

    _heading(doc, "验收与验证", 1)
    for text in [
        "未初始化 Owner 时无可登录默认账号，首次 Owner 只能在服务器本机 CLI 创建。",
        "未登录访问数据页/API 被拦截，Analyst 无法管理人员。",
        "重复终局快照不会重复产生对局或动物统计。",
        "网站不读取 player_accounts.json，也不输出密码哈希、设备 ID、令牌或原始账号库。",
        "默认绑定回环地址；远程监听缺少 TLS 时启动失败。",
        "Godot、Node、浏览器冒烟和窄屏可用性检查均通过。",
    ]:
        _bullet(doc, text)

    _heading(doc, "可编辑设计源", 1)
    _paragraph(doc, "Figma 设计文件：https://www.figma.com/design/bGtSRFlZfFe8erC5GdGVP4。该文件用于维护登录、总览、排行榜、玩家卡组、动物胜率和授权管理状态的 UI/UE 信息。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
